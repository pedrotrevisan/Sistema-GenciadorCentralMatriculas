from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Criar documento
doc = Document()

# Configurar estilos
style = doc.styles['Heading 1']
style.font.color.rgb = RGBColor(0, 69, 135)  # Azul SENAI
style.font.size = Pt(24)
style.font.bold = True

style2 = doc.styles['Heading 2']
style2.font.color.rgb = RGBColor(0, 69, 135)
style2.font.size = Pt(16)
style2.font.bold = True

# ==================== CAPA ====================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Sistema Central de Matrículas (CM)")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 69, 135)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("SENAI CIMATEC")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(227, 6, 19)  # Vermelho SENAI

doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run("Hub de Inteligência Operacional")
run.font.size = Pt(16)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Data
data_p = doc.add_paragraph()
data_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = data_p.add_run("Janeiro 2026")
run.font.size = Pt(12)

doc.add_page_break()

# ==================== VISÃO GERAL ====================
doc.add_heading("1. Visão Geral", level=1)

doc.add_paragraph(
    "O Sistema Central de Matrículas é uma plataforma web completa desenvolvida para "
    "digitalizar, automatizar e otimizar todo o processo de gestão de matrículas do "
    "SENAI CIMATEC, substituindo processos manuais e planilhas descentralizadas por "
    "uma solução integrada e inteligente."
)

doc.add_paragraph()

# ==================== PROBLEMAS RESOLVIDOS ====================
doc.add_heading("2. Problemas Resolvidos", level=1)

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'

# Cabeçalho
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Antes ❌"
hdr_cells[1].text = "Depois ✅"
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True

# Dados
dados = [
    ("Planilhas Excel dispersas", "Sistema centralizado único"),
    ("Dados duplicados e inconsistentes", "Validação automática de CPF"),
    ("Sem visibilidade do funil", "Dashboard analítico em tempo real"),
    ("Cadastro manual aluno por aluno", "Importação em lote via planilha"),
    ("Difícil rastrear status", "Fluxo de status automatizado"),
    ("Exportação manual para TOTVS", "Exportação automática formatada"),
]

for i, (antes, depois) in enumerate(dados, 1):
    table.rows[i].cells[0].text = antes
    table.rows[i].cells[1].text = depois

doc.add_paragraph()

# ==================== PERFIS DE ACESSO ====================
doc.add_heading("3. Perfis de Acesso (RBAC)", level=1)

doc.add_paragraph("O sistema possui 3 níveis de acesso com permissões específicas:")

doc.add_heading("Consultor", level=2)
doc.add_paragraph("• Cria novos pedidos de matrícula", style='List Bullet')
doc.add_paragraph("• Visualiza apenas seus próprios pedidos", style='List Bullet')
doc.add_paragraph("• Acompanha o status das suas solicitações", style='List Bullet')

doc.add_heading("Assistente", level=2)
doc.add_paragraph("• Visualiza todos os pedidos do sistema", style='List Bullet')
doc.add_paragraph("• Altera status dos pedidos (aprova, reprova, etc.)", style='List Bullet')
doc.add_paragraph("• Exporta dados para o TOTVS", style='List Bullet')
doc.add_paragraph("• Acessa dashboard completo", style='List Bullet')

doc.add_heading("Administrador", level=2)
doc.add_paragraph("• Todas as permissões anteriores", style='List Bullet')
doc.add_paragraph("• Gerencia usuários (criar, editar, desativar)", style='List Bullet')
doc.add_paragraph("• Gerencia cadastros (cursos, projetos, empresas)", style='List Bullet')
doc.add_paragraph("• Acesso total ao sistema", style='List Bullet')

doc.add_page_break()

# ==================== FUNCIONALIDADES ====================
doc.add_heading("4. Funcionalidades Implementadas", level=1)

# Dashboard
doc.add_heading("4.1 Dashboard Analítico 2.0", level=2)

doc.add_paragraph("Métricas em tempo real:", style='Intense Quote')
doc.add_paragraph("• Total de Pedidos - Visão geral do volume", style='List Bullet')
doc.add_paragraph("• Tempo Médio de Matrícula - Da criação à conclusão", style='List Bullet')
doc.add_paragraph("• Taxa de Conversão - Percentual de aprovações", style='List Bullet')
doc.add_paragraph("• Pedidos Críticos - Parados há mais de 48h (alerta)", style='List Bullet')

doc.add_paragraph()
doc.add_paragraph("Gráficos visuais:", style='Intense Quote')
doc.add_paragraph("• Funil de Matrículas - Visualize onde o processo está travando", style='List Bullet')
doc.add_paragraph("• Top 5 Empresas - Empresas com mais matrículas", style='List Bullet')
doc.add_paragraph("• Top 5 Projetos - Projetos mais demandados", style='List Bullet')
doc.add_paragraph("• Evolução Mensal - Tendência dos últimos 6 meses", style='List Bullet')

# Cadastro
doc.add_heading("4.2 Cadastro de Matrícula (Wizard 3 Etapas)", level=2)

doc.add_paragraph("Etapa 1 - Dados do Curso:", style='Intense Quote')
doc.add_paragraph("• Seleção do curso", style='List Bullet')
doc.add_paragraph("• Vínculo com Projeto ou Empresa", style='List Bullet')

doc.add_paragraph("Etapa 2 - Dados do Aluno:", style='Intense Quote')
doc.add_paragraph("• Informações pessoais completas", style='List Bullet')
doc.add_paragraph("• Validação automática de CPF (formato e dígitos)", style='List Bullet')
doc.add_paragraph("• Bloqueio de CPF duplicado no sistema", style='List Bullet')
doc.add_paragraph("• Dados de contato e endereço", style='List Bullet')

doc.add_paragraph("Etapa 3 - Revisão e Confirmação:", style='Intense Quote')
doc.add_paragraph("• Resumo visual de todos os dados", style='List Bullet')
doc.add_paragraph("• Confirmação antes de enviar", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("✨ Diferencial: ")
run.bold = True
p.add_run("Formatação automática de nomes")

doc.add_paragraph("• JOÃO CARLOS DA SILVA → João Carlos da Silva", style='List Bullet')
doc.add_paragraph("• MARIA DAS DORES → Maria das Dores", style='List Bullet')

# Importação em Lote
doc.add_heading("4.3 Importação em Lote (NOVO!)", level=2)

p = doc.add_paragraph()
run = p.add_run("Cadastre dezenas de alunos de uma só vez!")
run.bold = True
run.font.color.rgb = RGBColor(227, 6, 19)

doc.add_paragraph("Como funciona:", style='Intense Quote')
doc.add_paragraph("1. Baixe o template Excel padronizado", style='List Number')
doc.add_paragraph("2. Preencha os dados dos alunos na planilha", style='List Number')
doc.add_paragraph("3. Faça upload no sistema", style='List Number')
doc.add_paragraph("4. Validação automática de todos os campos", style='List Number')
doc.add_paragraph("5. Preview visual com erros destacados", style='List Number')
doc.add_paragraph("6. Importação com um clique", style='List Number')

doc.add_paragraph()
doc.add_paragraph("Validações automáticas:", style='Intense Quote')
doc.add_paragraph("• ✅ CPF válido e não duplicado", style='List Bullet')
doc.add_paragraph("• ✅ Email no formato correto", style='List Bullet')
doc.add_paragraph("• ✅ Telefone com DDD", style='List Bullet')
doc.add_paragraph("• ✅ Campos obrigatórios preenchidos", style='List Bullet')
doc.add_paragraph("• ✅ UF válida (2 caracteres)", style='List Bullet')

doc.add_page_break()

# Gestão de Pedidos
doc.add_heading("4.4 Gestão de Pedidos", level=2)

doc.add_paragraph("Listagem completa com:", style='Intense Quote')
doc.add_paragraph("• Filtros por status", style='List Bullet')
doc.add_paragraph("• Busca por nome/CPF", style='List Bullet')
doc.add_paragraph("• Paginação", style='List Bullet')
doc.add_paragraph("• Ordenação", style='List Bullet')

doc.add_paragraph()
doc.add_paragraph("Fluxo de Status:", style='Intense Quote')
p = doc.add_paragraph()
p.add_run("Pendente → Em Análise → Doc. Pendente → Aprovado → Realizado → Exportado").font.name = 'Consolas'

doc.add_paragraph()
doc.add_paragraph("Página de Detalhes:", style='Intense Quote')
doc.add_paragraph("• Todos os dados do pedido", style='List Bullet')
doc.add_paragraph("• Dados completos do aluno", style='List Bullet')
doc.add_paragraph("• Histórico de alterações", style='List Bullet')
doc.add_paragraph("• Ações rápidas (alterar status)", style='List Bullet')

# Exportação TOTVS
doc.add_heading("4.5 Exportação TOTVS", level=2)

doc.add_paragraph("Gere arquivos prontos para importação no TOTVS:")
doc.add_paragraph("• Formato XLSX (Excel)", style='List Bullet')
doc.add_paragraph("• Formato CSV", style='List Bullet')
doc.add_paragraph("• Apenas pedidos com status 'Realizado'", style='List Bullet')
doc.add_paragraph("• Campos formatados conforme padrão TOTVS", style='List Bullet')
doc.add_paragraph("• Marcação automática como 'Exportado'", style='List Bullet')

# Gestão de Cadastros
doc.add_heading("4.6 Gestão de Cadastros (Admin)", level=2)

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'

hdr = table2.rows[0].cells
hdr[0].text = "Cadastro"
hdr[1].text = "Funcionalidades"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

cadastros = [
    ("Cursos", "Criar, editar, desativar cursos"),
    ("Projetos", "Gerenciar projetos SENAI"),
    ("Empresas", "Cadastro de empresas parceiras"),
    ("Usuários", "Criar, editar, alterar permissões"),
]

for i, (cad, func) in enumerate(cadastros, 1):
    table2.rows[i].cells[0].text = cad
    table2.rows[i].cells[1].text = func

doc.add_page_break()

# ==================== SEGURANÇA ====================
doc.add_heading("5. Segurança e Qualidade", level=1)

doc.add_paragraph("• ✅ Autenticação JWT - Tokens seguros com expiração", style='List Bullet')
doc.add_paragraph("• ✅ Controle de Acesso (RBAC) - Permissões por perfil", style='List Bullet')
doc.add_paragraph("• ✅ Validação de Dados - Frontend e backend", style='List Bullet')
doc.add_paragraph("• ✅ Prevenção de Duplicidade - CPF único no sistema", style='List Bullet')
doc.add_paragraph("• ✅ Auditoria - Registro de todas as ações", style='List Bullet')

# ==================== BENEFÍCIOS ====================
doc.add_heading("6. Benefícios Esperados", level=1)

table3 = doc.add_table(rows=6, cols=2)
table3.style = 'Table Grid'

hdr = table3.rows[0].cells
hdr[0].text = "Métrica"
hdr[1].text = "Impacto"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

beneficios = [
    ("Tempo de cadastro", "⬇️ Redução de 70% com importação em lote"),
    ("Erros de dados", "⬇️ Redução de 90% com validações automáticas"),
    ("Visibilidade", "⬆️ Dashboard em tempo real para gestão"),
    ("Produtividade", "⬆️ Equipe focada em análise, não digitação"),
    ("Rastreabilidade", "⬆️ Histórico completo de cada pedido"),
]

for i, (metrica, impacto) in enumerate(beneficios, 1):
    table3.rows[i].cells[0].text = metrica
    table3.rows[i].cells[1].text = impacto

doc.add_paragraph()

# ==================== ROADMAP ====================
doc.add_heading("7. Próximas Evoluções (Roadmap)", level=1)

doc.add_heading("Curto Prazo (P1)", level=2)
doc.add_paragraph("• 📜 Timeline de Auditoria Visual - Histórico completo na tela", style='List Bullet')
doc.add_paragraph("• 🔔 Central de Pendências - Alertas de documentos faltantes", style='List Bullet')
doc.add_paragraph("• 📅 Filtros por Data - Período personalizado", style='List Bullet')

doc.add_heading("Médio Prazo (P2)", level=2)
doc.add_paragraph("• 💰 Controle Orçamentário - Gestão financeira por projeto", style='List Bullet')
doc.add_paragraph("• 📧 Notificações por Email - Alertas automáticos", style='List Bullet')
doc.add_paragraph("• 📍 Integração ViaCEP - Preenchimento automático de endereço", style='List Bullet')

doc.add_heading("Longo Prazo", level=2)
doc.add_paragraph("• 📊 Relatórios PDF - Geração automática", style='List Bullet')
doc.add_paragraph("• 🔗 Integração TOTVS Real-time - Sincronização automática", style='List Bullet')
doc.add_paragraph("• 📱 App Mobile - Acesso pelo celular", style='List Bullet')

doc.add_page_break()

# ==================== TECNOLOGIAS ====================
doc.add_heading("8. Tecnologias Utilizadas", level=1)

table4 = doc.add_table(rows=6, cols=2)
table4.style = 'Table Grid'

hdr = table4.rows[0].cells
hdr[0].text = "Camada"
hdr[1].text = "Tecnologia"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

techs = [
    ("Frontend", "React 19, Tailwind CSS, Shadcn/UI"),
    ("Backend", "Python, FastAPI, SQLAlchemy"),
    ("Banco de Dados", "PostgreSQL / SQLite"),
    ("Autenticação", "JWT (JSON Web Tokens)"),
    ("Arquitetura", "Clean Architecture + DDD"),
]

for i, (camada, tech) in enumerate(techs, 1):
    table4.rows[i].cells[0].text = camada
    table4.rows[i].cells[1].text = tech

doc.add_paragraph()

# ==================== CREDENCIAIS ====================
doc.add_heading("9. Demonstração", level=1)

doc.add_paragraph("Credenciais de Teste:")

table5 = doc.add_table(rows=4, cols=3)
table5.style = 'Table Grid'

hdr = table5.rows[0].cells
hdr[0].text = "Perfil"
hdr[1].text = "Email"
hdr[2].text = "Senha"
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

creds = [
    ("Admin", "admin@senai.br", "admin123"),
    ("Assistente", "assistente@senai.br", "assistente123"),
    ("Consultor", "consultor@senai.br", "consultor123"),
]

for i, (perfil, email, senha) in enumerate(creds, 1):
    table5.rows[i].cells[0].text = perfil
    table5.rows[i].cells[1].text = email
    table5.rows[i].cells[2].text = senha

doc.add_paragraph()

# ==================== RESUMO ====================
doc.add_heading("10. Resumo Executivo", level=1)

doc.add_paragraph("O Sistema Central de Matrículas é uma solução completa que:")

doc.add_paragraph("1. Centraliza toda a gestão de matrículas em um único lugar", style='List Number')
doc.add_paragraph("2. Automatiza validações e formatações de dados", style='List Number')
doc.add_paragraph("3. Agiliza o cadastro com importação em lote", style='List Number')
doc.add_paragraph("4. Fornece visibilidade com dashboard analítico", style='List Number')
doc.add_paragraph("5. Garante qualidade dos dados com validações rigorosas", style='List Number')
doc.add_paragraph("6. Integra com o TOTVS através de exportação formatada", style='List Number')
doc.add_paragraph("7. Escala para atender o crescimento da demanda", style='List Number')

doc.add_paragraph()
doc.add_paragraph()

# Rodapé
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Desenvolvido com 💙 para o SENAI CIMATEC")
run.italic = True
run.font.color.rgb = RGBColor(0, 69, 135)

# Salvar
doc.save('/app/Apresentacao_Sistema_Central_Matriculas.docx')
print("✅ Documento criado com sucesso!")
print("📄 Arquivo: /app/Apresentacao_Sistema_Central_Matriculas.docx")
